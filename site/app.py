import os
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_from_directory
from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document
from dropbox_uploader import upload_to_dropbox

app = Flask(__name__)
app.secret_key = 'supersecretkey'

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

USERS = {
    'Lucas': {'senha': '12345', 'role': 'admin'},
    'Andre': {'senha': '12345', 'role': 'admin'},
    'Erismar': {'senha': '0421', 'role': 'admin'},
    'Loja1': {'senha': '12345', 'role': 'admin'},
    'Josuel': {'senha': '1999', 'role': 'admin'},
    'Celso': {'senha': '030257', 'role': 'funcionario'},
    'Leidvan': {'senha': '12345', 'role': 'funcionario'},
    'Henrique': {'senha': '2403', 'role': 'funcionario'},
}

EQUIPAMENTOS = [
    'betoneira',
    'compactador de solo',
    'placa vibratória',
    'alisadora de piso',
    'gerador',
    'cortadora de piso',
    'martelete',
    'motosserra',
    'motobomba',
    'compressor',
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    if 'user' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = request.form['username']
        senha = request.form['password']
        if user in USERS and USERS[user]['senha'] == senha:
            session['user'] = user
            session['role'] = USERS[user]['role']
            flash(f'Bem vindo, {user}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Usuário ou senha inválidos', 'danger')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Desconectado com sucesso.', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user' not in session:
        return redirect(url_for('login'))
    return render_template('dashboard.html', user=session['user'], role=session['role'], equipamentos=EQUIPAMENTOS)

@app.route('/checklist/<equipamento>', methods=['GET', 'POST'])
def checklist(equipamento):
    if 'user' not in session:
        return redirect(url_for('login'))
    equipamento = equipamento.lower()
    if equipamento not in [e.lower() for e in EQUIPAMENTOS]:
        flash('Equipamento inválido.', 'danger')
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        descricao = request.form.get('descricao')
        fotos = request.files.getlist('fotos')
        if not descricao:
            flash('Descrição é obrigatória.', 'danger')
            return redirect(request.url)
        
        hoje = datetime.now().strftime('%Y-%m-%d')
        pasta = os.path.join(app.config['UPLOAD_FOLDER'], hoje, equipamento.replace(' ', '_'))
        os.makedirs(pasta, exist_ok=True)

        fotos_filenames = []
        for foto in fotos:
            if foto and allowed_file(foto.filename):
                filename = secure_filename(foto.filename)
                caminho = os.path.join(pasta, filename)
                foto.save(caminho)
                fotos_filenames.append(filename)
        
        doc = Document()
        doc.add_heading(f'Checklist - {equipamento.capitalize()}', 0)
        doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph(f'Funcionário: {session["user"]}')
        doc.add_paragraph('Descrição do serviço realizado:')
        doc.add_paragraph(descricao)
        if fotos_filenames:
            doc.add_paragraph('Fotos anexadas:')
            for f in fotos_filenames:
                doc.add_paragraph(f'- {f}')
        else:
            doc.add_paragraph('Nenhuma foto enviada.')

        nome_doc = f'checklist_{equipamento}_{datetime.now().strftime("%Y%m%d%H%M%S")}_{session["user"]}.docx'
        caminho_doc = os.path.join(pasta, nome_doc)
        doc.save(caminho_doc)

        # Envia ao Dropbox
        dropbox_path = f"/checklists/{nome_doc}"
        try:
            link = upload_to_dropbox(caminho_doc, dropbox_path)
            print(f"📁 Checklist enviado ao Dropbox com sucesso! Link: {link}")
        except Exception as e:
            print(f"⚠️ Erro ao enviar para o Dropbox: {e}")

        flash('Checklist salvo com sucesso!', 'success')
        return redirect(url_for('dashboard'))

    return render_template('checklist.html', equipamento=equipamento.capitalize())

@app.route('/admin')
def admin_panel():
    if 'user' not in session or session.get('role') != 'admin':
        flash('Acesso negado.', 'danger')
        return redirect(url_for('login'))
    
    base_dir = app.config['UPLOAD_FOLDER']
    filtro_equip = request.args.get('equipamento', '').lower()
    filtro_data = request.args.get('data', '')

    checklists = []
    if os.path.exists(base_dir):
        for data_dir in sorted(os.listdir(base_dir), reverse=True):
            if filtro_data and filtro_data != data_dir:
                continue
            caminho_data = os.path.join(base_dir, data_dir)
            if os.path.isdir(caminho_data):
                for equipamento_dir in os.listdir(caminho_data):
                    if filtro_equip and filtro_equip != equipamento_dir:
                        continue
                    caminho_equip = os.path.join(caminho_data, equipamento_dir)
                    if os.path.isdir(caminho_equip):
                        for arquivo in os.listdir(caminho_equip):
                            if arquivo.endswith('.docx'):
                                nome_operador = arquivo.split('_')[-1].replace('.docx', '')
                                fotos = [f for f in os.listdir(caminho_equip) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
                                try:
                                    nome_operador = arquivo.rsplit('_', 1)[-1].replace('.docx', '')
                                except:
                                    nome_operador = 'Desconhecido'
                                checklists.append({
                                    'data': data_dir,
                                    'equipamento': equipamento_dir.replace('_', ' '),
                                    'arquivo': arquivo,
                                    'caminho': f'{data_dir}/{equipamento_dir}/{arquivo}',
                                    'fotos': fotos[:3],
                                    'operador': nome_operador
                                })
    checklists = sorted(checklists, key=lambda x: x['data'], reverse=True)

    return render_template('admin_panel.html', checklists=checklists, equipamentos=EQUIPAMENTOS)

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
